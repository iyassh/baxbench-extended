from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

doc = Document()

def set_style(style, font_name, font_size, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    style.paragraph_format.alignment = align

# Setup Styles matching IEEE
set_style(doc.styles['Title'], 'Times New Roman', 24, align=WD_ALIGN_PARAGRAPH.CENTER)
set_style(doc.styles['Heading 1'], 'Times New Roman', 10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.styles['Heading 1'].font.small_caps = True
set_style(doc.styles['Heading 2'], 'Times New Roman', 10, italic=True)
set_style(doc.styles['Normal'], 'Times New Roman', 10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# Title & Authors (One column section)
title = doc.add_paragraph("CodeStrike Extended: Can AI Write Secure Web Applications?", style='Title')
authors = doc.add_paragraph("Yaash | Ravinder | Deepansh Sharma | Vansh\nDepartment of Computer Science\nCOMP 4210 Ethical Hacking\n", style='Normal')
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Create new section for Two Columns
new_section = doc.add_section()
sectPr = new_section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '2')
cols.set(qn('w:space'), '360') # 0.25 inch gap

with open('/Users/deepanshsharma/baxbench-extended/docs/research_report.md', 'r') as f:
    for line in f.readlines():
        line = line.strip()
        if not line: continue
        
        # Skip header metadata we already added
        if 'CodeStrike Extended' in line or 'Yaash' in line or 'COMP 4210' in line or 'Department' in line or line.startswith('---'):
            continue
            
        if line.startswith('![CodeStrike Security Dashboard]'):
            try:
                doc.add_picture('/Users/deepanshsharma/baxbench-extended/docs/dashboard.png', width=Inches(3.2))
            except:
                pass
            continue
            
        clean_text = line.replace('**', '').replace('_', '').replace('#', '').strip()
        
        if line.startswith('## Abstract') or line.startswith('## Keywords'):
            continue
        if line.startswith('Abstract') or line.startswith('Keywords'):
            p = doc.add_paragraph(clean_text, style='Normal')
            p.runs[0].bold = True
            p.runs[0].italic = True
            continue
            
        if line.startswith('## '):
            doc.add_paragraph(clean_text, style='Heading 1')
            continue
        if line.startswith('### '):
            doc.add_paragraph(clean_text, style='Heading 2')
            continue
            
        doc.add_paragraph(clean_text, style='Normal')

doc.save('/Users/deepanshsharma/baxbench-extended/docs/CodeStrike_Report_IEEE_Format.docx')
