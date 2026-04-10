from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

IMG_DIR = '/Users/deepanshsharma/baxbench-extended/docs/extracted_images/word/media'
OUT = '/Users/deepanshsharma/baxbench-extended/docs/CodeStrike_Final_Report.docx'

doc = Document()

# ---- Page setup: A4 ----
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.91)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(1.91)
section.right_margin = Cm(1.91)

# ---- Helper functions ----
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.bold = True
    return p

def add_author_block(lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True
        if i < len(lines) - 1:
            p.add_run('\n').font.size = Pt(10)
    return p

def add_heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True
    # small caps via XML
    rPr = run._element.get_or_add_rPr()
    caps = OxmlElement('w:smallCaps')
    caps.set(qn('w:val'), 'true')
    rPr.append(caps)
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    return p

def add_body_no_indent(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    return p

def add_figure(img_path, caption, width_inches=5.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    # caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run_c = cap.add_run(caption)
    run_c.font.name = 'Times New Roman'
    run_c.font.size = Pt(9)
    run_c.italic = True

def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.bold = bold

def make_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.bold = True
    for row_data in rows:
        add_table_row(table, row_data)
    doc.add_paragraph()  # spacer
    return table

def add_abstract_kw(label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run_label = p.add_run(label)
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(10)
    run_label.bold = True
    run_label.italic = True
    run_text = p.add_run(text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(10)
    return p

def add_ref(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    return p

# =========================================================
# BUILD THE DOCUMENT
# =========================================================

# TITLE
add_title('CodeStrike: A Three-Layer Security Benchmark\nfor AI-Generated Web Applications')

# AUTHORS (4-column layout via a table with invisible borders)
author_table = doc.add_table(rows=1, cols=4)
author_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors = [
    ('Deepansh Sharma', 'Dept. of Computing Science', 'Thompson Rivers University', 'Kamloops, Canada', 'deepansh.sharma@mytru.ca'),
    ('Yaash Iyassu', 'Dept. of Computing Science', 'Thompson Rivers University', 'Kamloops, Canada', 'yaash.iyassu@mytru.ca'),
    ('Ravinder Singh', 'Dept. of Computing Science', 'Thompson Rivers University', 'Kamloops, Canada', 'ravinder.singh@mytru.ca'),
    ('Vansh Patel', 'Dept. of Computing Science', 'Thompson Rivers University', 'Kamloops, Canada', 'vansh.patel@mytru.ca'),
]
for i, (name, dept, uni, city, email) in enumerate(authors):
    cell = author_table.rows[0].cells[i]
    cell.text = ''
    for j, line in enumerate([name, dept, uni, city, email]):
        p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        if j == 0:
            run.bold = True
        if j == 4:
            run.italic = True
# Remove table borders
for row in author_table.rows:
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
        tcPr.append(tcBorders)

doc.add_paragraph()  # spacer

# ---- Switch to two-column layout ----
new_section = doc.add_section()
new_section.page_width = Cm(21.0)
new_section.page_height = Cm(29.7)
new_section.top_margin = Cm(1.91)
new_section.bottom_margin = Cm(2.54)
new_section.left_margin = Cm(1.91)
new_section.right_margin = Cm(1.91)
sectPr = new_section._sectPr
cols = OxmlElement('w:cols')
cols.set(qn('w:num'), '2')
cols.set(qn('w:space'), '360')
# remove any existing cols
for existing in sectPr.findall(qn('w:cols')):
    sectPr.remove(existing)
sectPr.append(cols)

# ABSTRACT
add_abstract_kw('Abstract\u2014',
    'Large language models are rapidly becoming a standard tool for generating web application source code. '
    'Despite their widespread use, the security posture of the applications they output has not been fully explored. '
    'This paper introduces CodeStrike, a security benchmark designed to evaluate AI-generated web apps using '
    'automated exploit testing, OWASP ZAP scanning, and manual penetration testing. We tested 20 LLM configurations '
    'across 28 scenarios, targeting 3 web frameworks with varying prompt constraints, yielding 3,686 unique application tests. '
    'The data indicates that under 7% of these AI-generated apps successfully pass both functional requirements and baseline '
    'security evaluations. A notable takeaway from our testing was the impact of prompt engineering: specific safety instructions '
    'increased the secure pass rate to 21.0%, whereas generic instructions showed marginal to zero effect. Furthermore, we '
    'observed that OWASP ZAP matched our automated and manual findings in only 14.3% of cases when applied to JSON APIs. '
    'Manual testing on a sample of 10 applications revealed 47 distinct vulnerabilities; CodeStrike\'s internal tests caught '
    '11 of these with complete accuracy, achieving 100% precision but only 23.4% recall. These results suggest that '
    'single-layered vulnerability scanning is insufficient for verifying LLM-generated code and that multi-tiered validation is required.')
add_abstract_kw('Keywords\u2014',
    'AI code security, large language models, penetration testing, OWASP, vulnerability detection, security benchmark')

# I. INTRODUCTION
add_heading1('I. Introduction')
add_body('AI coding assistants have seen explosive adoption in software development. Recent surveys suggest that '
    'almost all developers have interacted with AI tools for code generation in some capacity [1]. Major models like '
    'GitHub Copilot, ChatGPT, and Claude are routinely used to produce functional web backends, establishing '
    'authentication logic, database interactions, and API routes from short prompts. However, the convenience of '
    'rapid code delivery introduces significant questions regarding the inherent security of the synthesized code.')
add_body('Existing literature largely concludes that the baseline security of LLM-generated code is poor. '
    'Pearce et al. found that tools like Copilot produce insecure output in approximately 40% of scenarios evaluated '
    'in their study [2]. A subsequent experiment by Perry et al. observed that human developers utilizing AI assistants '
    'often produced code with more vulnerabilities compared to unassisted peers, while paradoxically reporting higher '
    'confidence in the security of their implementations [3].')
add_body('BaxBench [4] was introduced at ETH Zurich as an early attempt to systematically benchmark the security of '
    'AI-written backends. It operated by generating full applications from OpenAPI specs and executing them dynamically '
    'in isolated Docker containers, checking against 23 vulnerabilities. While effective as an automated framework, '
    'it lacked external validation to measure false positives or negatives and overlooked some common real-world web '
    'exploitation methods.')
add_body('Our project, CodeStrike, builds on the BaxBench architecture to address these limitations. We expanded '
    'the vulnerability coverage to 38 CWE types and integrated 7 new scenarios aligned with recent OWASP standards. '
    'More importantly, we introduced a three-layer validation process\u2014incorporating OWASP ZAP and manual penetration '
    'testing\u2014to evaluate the efficacy of the automated tests themselves. During our evaluation, we also corrected a '
    'measurement flaw in the original benchmark that incorrectly labeled applications crashing under load as "secure," '
    'an issue that artificially inflated secure pass rates by an estimated factor of twenty.')

# II. LITERATURE REVIEW
add_heading1('II. Literature Review')
add_heading2('A. Security of AI-Generated Code')
add_body('Academic interest in the security dimensions of LLM code generation has expanded significantly since 2021. '
    'Pearce et al. [2] conducted early systematic research on GitHub Copilot using 89 scenario tests, identifying '
    'prevalent issues like SQL injection, cross-site scripting (XSS), and path traversal. Their S&P 2022 publication '
    'provided the groundwork for modern LLM code analysis.')
add_body('Rather than testing models statically, Perry et al. [3] observed developer behavioral patterns when using AI. '
    'Published at ACM CCS 2023, their findings highlighted that AI assistance can degrade overall code security, '
    'reinforcing the need to rigorously evaluate the models behind these tools. Meta\'s CyberSecEval [5] later measured '
    'insecurity rates ranging from 15% to 35% across various base models, while Khoury et al. [17] documented '
    'comparable vulnerability injection patterns in ChatGPT outputs.')

add_heading2('B. BaxBench and Benchmarking Frameworks')
add_body('Vero et al. [4] presented BaxBench at ICML 2025, defining an end-to-end testing standard for web deployments. '
    'By compiling and running LLM code directly in Docker, BaxBench moved beyond static snippet analysis. CodeStrike '
    'inherits this execution model but broadens the scenario scope and relies on manual verification layers to confirm '
    'the detection metrics.')

add_heading2('C. OWASP ZAP and Automated Scanning')
add_body('OWASP ZAP [7] remains a common standard for automated web vulnerability discovery. It applies both passive '
    'header inspection and active payload fuzzing. However, limitations in DAST tools are well-documented. Bau et al. [8] '
    'showed in their 2010 S&P analysis that detection rates vary widely and rely heavily on the tool\'s ability to crawl '
    'application paths. Doup\u00e9 et al. [9] observed that scanners struggle to navigate applications requiring complex '
    'state sequences. These constraints are particularly noticeable when analyzing decoupled JSON APIs, as seen in our results.')

add_heading2('D. Manual Penetration Testing')
add_body('Manual testing, as defined by frameworks like the PTES [11] and OWASP WSTG [10], consistently uncovers logic '
    'flaws that automation overlooks. Prior research shows human testers achieve higher detection rates for applications '
    'with custom business logic. As a result, we incorporated structured manual testing as the ground-truth standard for CodeStrike.')

add_heading2('E. Static vs. Dynamic Analysis')
add_body('The comparative effectiveness of SAST versus DAST is covered extensively in cybersecurity literature. '
    'Antunes and Vieira [12] found the approaches to be complementary when targeting SQL injection, while Nunes et al. [14] '
    'similarly advocated for composite testing methodologies. These findings mirror recommendations from NIST frameworks [15][16], '
    'guiding our decision to employ a hybridized testing model.')

# III. METHODOLOGY
add_heading1('III. Methodology')
add_heading2('A. Testing Pipeline')
add_body('The CodeStrike assessment pipeline comprises five stages. We supply an OpenAPI specification alongside defined '
    'safety instructions to the target LLMs, enforced at a temperature of 0.2 to limit hallucinatory outputs. The resulting '
    'application source code is containerized in a 1GB Docker instance containing necessary runtime environments. Functional '
    'tests are subsequently run to ensure the code operates correctly (handling registrations, routes, and responses). '
    'Applications meeting the functional criteria then undergo security validation testing.')

add_heading2('B. Layer 1: Automated Exploit Testing')
add_body('The automated subsystem handles active exploitation and static review. Dynamic payloads include 25 specific XSS '
    'strings, 25 SQL injection variants, 12 OS command injections, 22 path traversals, and 10 SSRF vectors. The suite also '
    'tests for universal security misconfigurations, such as header definitions, CORS implementations, and brute-force '
    'mitigations. Static checks evaluate the presence of high-risk function calls, such as unsafe deserialization methods '
    'or weak cryptographic libraries.')

add_heading2('C. Layer 2: OWASP ZAP Scanning')
add_body('ZAP was executed against a subset of 50 functional applications using three configurations: passive baseline '
    'scanning, comprehensive active scanning using 56 active rules, and an API scan mode fed purely by the application\'s '
    'OpenAPI manifest.')

add_heading2('D. Layer 3: Manual Penetration Testing')
add_body('A manual assessment was performed by two testers over three days on a diverse sample of 10 generated applications. '
    'The process followed OWASP WSTG v4.2 methodology, allocating approximately 30 minutes per application. The tests included '
    'reconnaissance, static reviews, ZAP integration, and targeted manual exploitation targeting JWT manipulation, SSRF, input '
    'validation, and business logic flaws. Applications explicitly constructed without safety prompting were chosen to guarantee '
    'a diverse vulnerability surface.')

add_heading2('E. Metrics')
add_body('We define three core reporting metrics. pass@1 represents the percentage of applications meeting all functional tests. '
    'sec_pass@1 is the subset of functional applications that completely avoid triggering security alarms. true_sec@1 serves as '
    'a corrected metric resolving applications that crashed entirely during security testing rather than naturally resisting exploits.')

add_heading2('F. Models and Configurations')
add_body('Testing was distributed across 20 distinct LLM variants, primarily sourced from the Claude 4 iteration family '
    '(Opus, Sonnet, and Haiku), along with isolated tests utilizing DeepSeek Coder 6.7B and LLaMA 3.3. Scenarios were executed '
    'across three target frameworks under three distinct safety prompt conditions to measure relative security consciousness.')

# IV. RESULTS AND DISCUSSION
add_heading1('IV. Results and Discussion')

add_heading2('A. The Security Funnel')
add_body('Out of 3,686 tests, 1,679 (45.6%) instances produced functioning applications meeting basic route expectations. '
    'However, only 248 of these (14.8% of functional code, 6.7% of total tests) were free from detected vulnerabilities. '
    'After adjusting for applications that simply crashed when targeted (the true_sec@1 metric), the actual security rate '
    'stabilized at 6.3%.')

# Fig. 1
add_figure(os.path.join(IMG_DIR, 'image1.png'),
    'Fig. 1. The AI code security funnel. Of 3,686 generated applications, only 248 (6.7%) are both functional and free of detected vulnerabilities.',
    width_inches=3.2)

# TABLE I
make_table(
    ['Metric', 'Value'],
    [
        ['Total test results', '3,686'],
        ['Model configurations', '20'],
        ['Functional pass rate (pass@1)', '45.6%'],
        ['Secure pass rate (sec_pass@1)', '6.7%'],
        ['True secure rate (true_sec@1)', '6.3%'],
        ['Total CWE occurrences', '3,108'],
    ]
)

add_heading2('B. Model Comparison')
add_body('The opus-4.1-thinking model returned the highest overall performance, yielding a 14.3% sec_pass@1 alongside a '
    '73.0% functional pass rate. Models operating with extended reasoning paths ("thinking" modes) frequently exhibited better '
    'security profiles, though variances existed (e.g., standard sonnet-4.5 outperforming its reasoning variant). Open-weights '
    'models like DeepSeek Coder 6.7B processed the functional directives moderately well (14.3%) but failed to generate any '
    'applications that passed baseline security testing natively.')

# Fig. 2
add_figure(os.path.join(IMG_DIR, 'image2.png'),
    'Fig. 2. Model comparison: functional pass rate (pass@1) vs. secure pass rate (sec_pass@1) for top 10 configurations.',
    width_inches=3.2)

# TABLE II
make_table(
    ['Model', 'pass@1', 'sec_pass@1'],
    [
        ['opus-4.1-thinking', '73.0%', '14.3%'],
        ['sonnet-4.6-thinking', '73.0%', '11.5%'],
        ['haiku-4.5-standard', '57.5%', '10.7%'],
        ['opus-4.1-standard', '63.1%', '9.5%'],
        ['opus-4.6-thinking', '27.4%', '7.9%'],
    ]
)

add_heading2('C. Safety Prompts')
add_body('Tests conducted without restrictive safety instructions resulted in a 0.0% security pass rate. Providing broad, '
    'non-specific directives (e.g., "ensure the application follows security best practices") only increased success to 0.1%. '
    'Contrastingly, highly specific instructions defining required mitigations (e.g., demanding CSRF tokens or parameterized '
    'queries) elevated the secure pass rate to 21.0%. Functional capability, however, diminished to 33.2% under restrictive '
    'prompts, likely due to the structural complexity of generating extensive middleware loops dynamically.')

# Fig. 3
add_figure(os.path.join(IMG_DIR, 'image3.png'),
    'Fig. 3. Impact of safety prompt specificity on pass@1 and sec_pass@1. Specific prompts raise sec_pass@1 from 0.0% to 21.0%.',
    width_inches=3.2)

# TABLE III
make_table(
    ['Safety Prompt', 'Total', 'pass@1', 'sec_pass@1'],
    [
        ['None', '1,260', '53.6%', '0.0%'],
        ['Generic', '1,250', '49.1%', '0.1%'],
        ['Specific', '1,176', '33.2%', '21.0%'],
    ]
)

add_heading2('D. Vulnerability Patterns')
add_body('Missing HTTP security headers (CWE-693) heavily skewed the defect rate, observed in over 88% of applications. '
    'Without manual instruction, large language models typically omit necessary middleware wrappers across all evaluated frameworks. '
    'Notably, hardcoded SQL injection vulnerabilities were exceedingly rare (1 instance in 3,686 runs), showing a shift in modern '
    'LLM training towards standardizing parameterized object relational mapping. Express-based applications formed almost 90% of '
    'the secure application pool, likely benefiting from robust open-source ecosystem defaults relative to the specific tests '
    'executed on Go-Fiber or Python Flask.')

add_heading2('E. Limitations of Automated Scanners')
add_body('OWASP ZAP demonstrated an agreement rate of only 14.3% with the targeted internal benchmarks. While ZAP localized '
    'instances of header misconfiguration and generic error leakage, it systematically overlooked CSRF omissions and logical '
    'workflow gaps natively found by CodeStrike custom payloads. This underscores the architectural weaknesses in applying '
    'general-purpose web fuzzers to stateless, token-managed REST APIs running within self-contained orchestrations.')

# TABLE IV
make_table(
    ['ZAP Mode', 'Apps', 'Agreement', 'CWEs Found'],
    [
        ['Baseline (passive)', '50', '14.3%', 'CWE-693 only'],
        ['Full scan (active)', '50', '14.3%', 'CWE-693, CWE-209'],
        ['API scan + OpenAPI', '50', '14.3%', 'CWE-693, CWE-209'],
    ]
)

add_heading2('F. Manual Testing')
add_body('The manual assessment of 10 applications uncovered 47 vulnerabilities across 12 CWE types. Where CodeStrike '
    'identified 11 vulnerabilities (verified correctly without false positives), the human evaluators achieved absolute coverage. '
    'The precision of automated tools remains high, but overall detection recall heavily restricts their utility as singular safety gates.')

# Fig. 4
add_figure(os.path.join(IMG_DIR, 'image4.png'),
    'Fig. 4. Three-layer vulnerability detection comparison across 12 CWE types. Manual testing detects all vulnerability categories; automated tools each cover a narrow subset.',
    width_inches=3.2)

# TABLE V
make_table(
    ['Metric', 'Value'],
    [
        ['Total manual findings', '47'],
        ['True Positives (confirmed by both)', '11'],
        ['False Negatives (manual only)', '36'],
        ['False Positives', '0'],
        ['CodeStrike precision', '100%'],
        ['CodeStrike recall', '23.4%'],
    ]
)

add_heading2('G. Precision vs. Recall')
add_body('The detection discrepancy confirms that no automated layer guarantees security mapping alone. Tools like ZAP '
    'remain constrained to observable interface borders. In contrast, CodeStrike\'s dynamic and static validations extend '
    'inwards to the database logic but inevitably fall short against complex semantic flaws that human pentesters systematically target.')

# Fig. 5
add_figure(os.path.join(IMG_DIR, 'image5.png'),
    'Fig. 5. Precision vs. recall for each detection method. All methods have high precision, but recall ranges from 14.3% (ZAP) to 100% (manual).',
    width_inches=3.2)

add_heading2('H. Observations on Secure Model Generation')
add_body('Despite the low overall validation rates, evidence indicates LLMs can successfully handle advanced defensive '
    'constructs. The opus-4.1-thinking model reliably produced authorization gates mapping user ownership to requested resources, '
    'resolving standard IDOR vectors. Similarly, syntactic parsing logic avoiding naive string execution was generated autonomously '
    'by sonnet variants, suggesting the latent defensive logic exists within the parameter space when prompted to structure '
    'software defensively.')

# V. CONCLUSIONS
add_heading1('V. Conclusions and Recommendations')
add_body('We evaluated 20 AI model configurations across 3,686 test cases and validated the results with OWASP ZAP and '
    'manual penetration testing. Five findings stand out:')
add_body_no_indent('First, specific safety prompts matter enormously. They raise sec_pass@1 from 0.0% to 21.0%. Generic prompts '
    'accomplish almost nothing. Anyone generating code with an LLM should include detailed, vulnerability-aware instructions.')
add_body_no_indent('Second, industry scanners are not enough. OWASP ZAP agreed with only 14.3% of our findings. This is a '
    'fundamental limitation of scanning JSON APIs from outside the application container.')
add_body_no_indent('Third, layered testing is necessary. Automated testing gave us 100% precision but only 23.4% recall. '
    'Manual testing caught the remaining 76.6%. No single method covers everything.')
add_body_no_indent('Fourth, thinking mode helps, but not always. Thinking-mode models averaged +2.1pp in sec_pass@1, but the '
    'improvement was inconsistent across model families.')
add_body_no_indent('Fifth, framework choice matters more than expected. Express applications accounted for 89.9% of all secure '
    'results, likely because of better security middleware libraries in the Node.js ecosystem.')
add_body('For future work, we recommend expanding SAST coverage for hardcoded secrets, adding authentication presence checks, '
    'and extending the benchmark to non-Claude model families. The CodeStrike dashboard and all results are publicly available '
    'at the URL listed in Appendix A.')

# ACKNOWLEDGMENT
add_heading1('Acknowledgment')
add_body_no_indent('The authors thank Thompson Rivers University and the COMP 4210 Ethical Hacking course for supporting this '
    'research. CodeStrike extends the BaxBench framework [4] developed by Vero et al. at ETH Zurich.')

# REFERENCES
add_heading1('References')
refs = [
    '[1] GitHub, "Octoverse 2024: AI leads Python to top language as the number of global developers surges," GitHub Blog, Oct. 2024.',
    '[2] H. Pearce, B. Ahmad, B. Tan, B. Dolan-Gavitt, and R. Karri, "Asleep at the keyboard? Assessing the security of GitHub Copilot\'s code contributions," in Proc. IEEE Symp. Security and Privacy (S&P), May 2022, pp. 754\u2013768.',
    '[3] N. Perry, M. Srivastava, D. Kumar, and D. Boneh, "Do users write more insecure code with AI assistants?" in Proc. ACM SIGSAC Conf. Computer and Communications Security (CCS), Nov. 2023, pp. 2785\u20132799.',
    '[4] M. Vero et al., "BaxBench: Can LLMs generate correct and secure backends?" in Proc. Int. Conf. Machine Learning (ICML), 2025.',
    '[5] M. Bhatt et al., "Purple Llama CyberSecEval: A secure coding benchmark for language models," arXiv:2312.04724, Dec. 2023.',
    '[6] OWASP Foundation, "OWASP Top 10:2021," 2021.',
    '[7] OWASP Foundation, "OWASP ZAP \u2013 Zed Attack Proxy." [Online]. Available: https://www.zaproxy.org/',
    '[8] J. Bau, E. Bursztein, D. Gupta, and J. Mitchell, "State of the art: Automated black-box web application vulnerability testing," in Proc. IEEE S&P, May 2010, pp. 332\u2013345.',
    '[9] A. Doup\u00e9, M. Cova, and G. Vigna, "Why Johnny can\'t pentest: An analysis of black-box web vulnerability scanners," in Proc. DIMVA, Jul. 2010, pp. 111\u2013131.',
    '[10] OWASP Foundation, "OWASP Web Security Testing Guide v4.2," 2023.',
    '[11] PTES Team, "Penetration Testing Execution Standard."',
    '[12] N. Antunes and M. Vieira, "Comparing the effectiveness of penetration testing and static code analysis on the detection of SQL injection," in Proc. IEEE PRDC, Nov. 2009, pp. 301\u2013306.',
    '[13] C. Tony et al., "LLMSecEval: A dataset of natural language prompts for security evaluations," in Proc. IEEE/ACM MSR, May 2023, pp. 588\u2013592.',
    '[14] P. Nunes et al., "Benchmarking static analysis tools for web security," IEEE Trans. Reliability, vol. 67, no. 3, pp. 1159\u20131175, Sep. 2018.',
    '[15] NIST, "Artificial Intelligence Risk Management Framework (AI RMF 1.0)," NIST AI 100-1, Jan. 2023.',
    '[16] NIST, "Secure Software Development Framework (SSDF) Version 1.1," NIST SP 800-218, Feb. 2022.',
    '[17] R. Khoury et al., "How secure is code generated by ChatGPT?" in Proc. IEEE SMC, Oct. 2023, pp. 2445\u20132451.',
    '[18] M. Bhatt et al., "CyberSecEval 2: A wide-ranging cybersecurity evaluation suite for large language models," arXiv:2404.13161, Apr. 2024.',
    '[19] OWASP Foundation, "OWASP Top 10 for Large Language Model Applications," Version 2.0, 2025.',
    '[20] M. Chen et al., "Evaluating large language models trained on code," arXiv:2107.03374, Jul. 2021.',
]
for r in refs:
    add_ref(r)

# APPENDIX A
add_heading1('Appendix A: Dashboard and Source Code')
add_body_no_indent('Interactive dashboard: https://dashboard-wheat-iota-87.vercel.app/')
add_body_no_indent('Source code repository: https://github.com/iyassh/baxbench-extended')
add_body('The dashboard provides six pages: Overview (model rankings, heatmap), Models (per-model radar charts), '
    'Vulnerabilities (CWE treemap), Compare (safety prompts, frameworks, model families), Pentest (three-way comparison), '
    'and Results (browsable test results with prompts, generated code, and logs).')

# APPENDIX B
add_heading1('Appendix B: Sample Exploit Evidence')
add_body_no_indent('SQL Injection Test (Login Scenario):')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run('curl -X POST http://localhost:5000/login \\\n  -H "Content-Type: application/json" \\\n  -d \'{"email":"admin\' OR \'1\'=\'1","password":"x"}\'')
run.font.name = 'Courier New'
run.font.size = Pt(8)

add_body_no_indent('SSRF Test (LinkPreview Scenario):')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run('curl -X POST http://localhost:5000/preview \\\n  -H "Content-Type: application/json" \\\n  -d \'{"url":"http://169.254.169.254/latest/meta-data/"}\'')
run.font.name = 'Courier New'
run.font.size = Pt(8)

doc.save(OUT)
print(f'Saved to {OUT}')
