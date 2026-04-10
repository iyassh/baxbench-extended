from docx import Document
from docx.shared import Inches
import os

doc_path = '/Users/deepanshsharma/baxbench-extended/docs/Conference-template-A4.docx'
doc = Document(doc_path)

# Clear contents
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)

with open('/Users/deepanshsharma/baxbench-extended/docs/research_report.md', 'r') as f:
    for line in f.readlines():
        line = line.strip()
        if not line: continue
        
        if line.startswith('![CodeStrike Security Dashboard]'):
            try:
                doc.add_picture('/Users/deepanshsharma/baxbench-extended/docs/dashboard.png', width=Inches(6.0))
            except:
                pass
            continue
            
        clean_text = line.replace('**', '').replace('_', '').replace('#', '').strip()
        doc.add_paragraph(clean_text)

try:
    doc.save('/Users/deepanshsharma/baxbench-extended/docs/CodeStrike_Report_Formatted.docx')
except Exception as e:
    print(e)
